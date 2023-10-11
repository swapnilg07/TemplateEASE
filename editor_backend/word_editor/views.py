import json
from django.http import HttpResponse, JsonResponse
from docx import Document
from docx.shared import Pt
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def extract_hyperlinks(run):
    hyperlinks = []
    for fld in run._element.findall('.//w:hyperlink', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        href = fld.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}href")
        text = fld.text
        if href and text:
            hyperlinks.append((text, href))
    return hyperlinks

def extract_email_hyperlinks(run):
    text = run.text
    email_links = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    return email_links

def apply_format(run, styles):
    for style in styles:
        if style == 'BOLD':
            run.bold = True
        if style == 'ITALIC':
            run.italic = True
        if style == 'UNDERLINE':
            run.underline = True

def getStylingByChar(length,inlineStyleRanges):
    charStyle = [[] for _ in range(length)]
    for i in range(length):
        for inline_style in inlineStyleRanges:
            start = inline_style['offset']
            end = start + inline_style['length']
            if start <= i and i< end:
                charStyle[i].append(inline_style['style'])
    return charStyle



def parse_to_docx(request):
    data = json.loads(request.body.decode('utf-8'))
    html_content = data.get("content")
    doc = Document()

    for block in html_content['blocks']:
        text = block['text']
        paragraph = doc.add_paragraph()
        if 'inlineStyleRanges' in block and block['inlineStyleRanges']:
            inlineStyleRanges = block['inlineStyleRanges']
            charStyle = getStylingByChar(len(text),inlineStyleRanges)
            for i in range(len(charStyle)):
                run = paragraph.add_run(text[i:i+1])
                apply_format(run, charStyle[i])
        else:
            paragraph.add_run(text)

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save('EditedWordFile.docx')

    with open('EditedWordFile.docx', 'rb') as docx_file:
        response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=EditedWordFile.docx'
        return response

def parse_docx_to_html(request):
    if request.method == 'POST':
        uploaded_file = request.FILES['file']
        doc = Document(uploaded_file)
        html_output = []

        for paragraph in doc.paragraphs:
            html_paragraph = "<p>"
            for run in paragraph.runs:
                font_size = run.font.size
                if font_size:
                    html_paragraph += f'<span style="font-size:{font_size.pt}px;">'

                if run.bold:
                    html_paragraph += "<strong>"
                if run.italic:
                    html_paragraph += "<em>"

                if run.underline:
                    html_paragraph += "<u>"
                
                if run.underline:
                    html_paragraph += "</u>"
                if run.italic:
                    html_paragraph += "</em>"
                if run.bold:
                    html_paragraph += "</strong>"
                if font_size:
                    html_paragraph += "</span>"
                    
                email_links = extract_email_hyperlinks(run)
                
                if email_links:
                    for email in email_links:
                        email_html = f'<a href="mailto:{email}">{email}</a>'
                        text = run.text.replace(email, email_html)
                        run.text = text

                html_paragraph += run.text

            html_paragraph += "</p>"
            html_output.append(html_paragraph)

        text = "\n".join(html_output)
        html_content = '<div>' + text.replace('\n', '<br>') + '</div>'
        return JsonResponse({'html_content': html_content})
    return JsonResponse({'error': 'Invalid request method'}, status=400)
